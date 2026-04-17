"""
deutsch_deid.processors.doc_processor
--------------------------------------
Document reading and PII processing pipelines for file-based input.

Reads .pdf, .docx, and .txt files, extracts their plain text, then
delegates to text_processor for analysis and guarding.

Consumed by the public namespace objects in deutsch_deid/__init__.py::

    from deutsch_deid import analyze, guard

    analyze.doc("/path/to/file.pdf")                     # -> list[dict]
    guard.doc("/path/to/file.docx", config={"mode": "tag"})  # -> dict

Functions in this module
------------------------
    analyze(path, config) -> list[dict]   detect PII in a document file
    guard(path, config)   -> dict         detect + anonymize a document file
    read(path)            -> str          extract plain text only (no PII pipeline)

Exceptions
----------
    UnsupportedFormatError   raised for any extension not in SUPPORTED_EXTENSIONS
    FileNotFoundError        raised when the path does not exist (after the extension check)
"""

import os
import re
from typing import Dict, List, Optional, Set

from deutsch_deid.processors.text_processor import analyze as _analyze, guard as _guard

SUPPORTED_EXTENSIONS: Set[str] = {".pdf", ".docx", ".txt"}


class UnsupportedFormatError(ValueError):
    """Raised when a file with an unsupported extension is provided."""


# ---------------------------------------------------------------------------
# File reader
# ---------------------------------------------------------------------------

def read(path: str) -> str:
    """
    Extract plain text from a document file.

    Parameters
    ----------
    path : str
        Absolute or relative path to the file.
        Supported extensions: ``.pdf``, ``.docx``, ``.txt``

    Returns
    -------
    str
        Plain text content of the document.

    Raises
    ------
    UnsupportedFormatError
        If the file extension is not in SUPPORTED_EXTENSIONS.
    FileNotFoundError
        If the file does not exist.
    """
    ext = os.path.splitext(path)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFormatError(
            f"Unsupported format {ext!r}. Supported formats are: {supported}"
        )

    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path!r}")

    if ext == ".txt":
        return _read_txt(path)
    elif ext == ".pdf":
        return _read_pdf(path)
    else:
        return _read_docx(path)


# ---------------------------------------------------------------------------
# Processing pipelines
# ---------------------------------------------------------------------------

def analyze(path: str, config: Optional[Dict] = None) -> List[Dict]:
    """
    Read a document and return a list of PII findings.

    Parameters
    ----------
    path   : Path to .pdf, .docx, or .txt file.
    config : Optional detection config (same as text_processor.analyze).

    Returns
    -------
    list[dict]
        Each dict: ``{"type": str, "start": int, "end": int, "score": float}``

    Raises
    ------
    UnsupportedFormatError
        If the file extension is not .pdf, .docx, or .txt.
    FileNotFoundError
        If *path* does not exist.
    """
    return _analyze(read(path), config=config)


def guard(path: str, config: Optional[Dict] = None) -> Dict:
    """
    Read a document, anonymize its PII, and return the guarded result.

    Parameters
    ----------
    path   : Path to .pdf, .docx, or .txt file.
    config : Optional processing config (same as text_processor.guard).

    Returns
    -------
    dict
        ``guarded_text`` – text with PII replaced.
        ``findings``     – list of finding dicts.

    Raises
    ------
    UnsupportedFormatError
        If the file extension is not .pdf, .docx, or .txt.
    FileNotFoundError
        If *path* does not exist.
    """
    return _guard(read(path), config=config)


# ---------------------------------------------------------------------------
# Format-specific readers
# ---------------------------------------------------------------------------

def _read_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as fh:
            return fh.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="cp1252", errors="replace") as fh:
            return fh.read()
    except PermissionError:
        raise PermissionError(f"Permission denied: {path!r}")
    except IsADirectoryError:
        raise IsADirectoryError(f"Is a directory, not a file: {path!r}")


def _normalize_pdf_text(text: str) -> str:
    """
    Repair the two most common pypdf extraction artefacts:

    1. Word-per-line scattering — when pypdf emits each word on its own line
       separated by a whitespace-only line:
           ``aanhoudende\\n \\npijn\\n \\nop``
       These are collapsed back into a single space.

    2. Double (or triple) spacing — pypdf inserts explicit character spacing
       that results in ``Medisch  Verslag:`` instead of ``Medisch Verslag:``.
       All runs of 2+ spaces/tabs are reduced to one space.

    Genuine paragraph breaks (blank lines) are preserved up to one blank line.
    """

    text = re.sub(r"\n[ \t]+\n", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()


def _read_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            "pypdf is required for PDF support. "
            "Install it with:  pip install pypdf"
        ) from exc

    reader = PdfReader(path)
    raw = "\n".join(page.extract_text() or "" for page in reader.pages)
    return _normalize_pdf_text(raw)


def _read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX support. "
            "Install it with:  pip install python-docx"
        ) from exc

    doc = Document(path)
    parts: List[str] = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)
